import os
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src" / "packagee"
OUT_MD = ROOT / os.environ.get("OSPEDALE_DOC_MD", "DOCUMENTO_CODIGO_COMPLETO.md")
OUT_DOCX = ROOT / os.environ.get("OSPEDALE_DOCX", "DOCUMENTO_CODIGO_COMPLETO.docx")


def read(name):
    return (SRC / name).read_text(encoding="utf-8")


def source_stats(name):
    text = read(name)
    lines = text.splitlines()
    fields = []
    methods = []
    for line in lines:
        stripped = line.strip()
        if re.match(r"^(private|protected|public)\s+(final\s+)?[\w<>\[\]]+\s+\w+\s*(=|;)", stripped):
            fields.append(stripped)
        if re.match(r"^(public|private|protected)\s+.*\)\s*\{", stripped):
            methods.append(stripped)
    return len(lines), fields, methods


def add_md(lines, title, body=None, level=2):
    lines.append("#" * level + " " + title)
    if body:
        lines.append("")
        if isinstance(body, list):
            lines.extend(body)
        else:
            lines.append(body)
    lines.append("")


def bullet(lines, items):
    for item in items:
        lines.append(f"- {item}")
    lines.append("")


def class_section(lines, file_name, role, details, flows=None, notes=None):
    count, fields, methods = source_stats(file_name)
    add_md(lines, file_name, level=2)
    lines.append(f"**Ubicacion:** `src/packagee/{file_name}`")
    lines.append("")
    lines.append(f"**Tamano aproximado:** {count} lineas.")
    lines.append("")
    lines.append(f"**Responsabilidad principal:** {role}")
    lines.append("")
    add_md(lines, "Explicacion detallada", details, level=3)
    if flows:
        add_md(lines, "Flujo interno", flows, level=3)
    if fields:
        add_md(lines, "Campos importantes detectados en el codigo", level=3)
        for field in fields[:25]:
            lines.append(f"- `{field}`")
        if len(fields) > 25:
            lines.append(f"- ... y {len(fields) - 25} campos adicionales, principalmente componentes Swing generados por NetBeans.")
        lines.append("")
    if methods:
        add_md(lines, "Metodos y constructores importantes", level=3)
        for method in methods[:60]:
            lines.append(f"- `{method}`")
        if len(methods) > 60:
            lines.append(f"- ... y {len(methods) - 60} metodos/eventos adicionales generados o conectados por NetBeans.")
        lines.append("")
    if notes:
        add_md(lines, "Notas para sustentacion", notes, level=3)


def build_markdown():
    lines = []
    add_md(lines, "DOCUMENTO TECNICO COMPLETO DEL CODIGO - OSPEDALE", [
        "**Proyecto:** Ospedale - Parcial 3",
        "",
        "**NRC:** 2039",
        "",
        "**Integrantes:** Santiago Andres Moreno Rivera, Jhon Sebastian Afanador Rueda, Adrian Jose Martinez Matinez",
        "",
        "Este documento explica el codigo del proyecto clase por clase y frame por frame. No esta escrito como una respuesta corta para sustentar; esta escrito como un manual de estudio para entender que hace cada archivo, por que existe y como se conecta con los demas.",
    ], level=1)

    add_md(lines, "1. Vision general del proyecto", [
        "Ospedale es una aplicacion de escritorio Java Swing construida en NetBeans. El proyecto fue reorganizado hacia una arquitectura MVC: los modelos representan los datos del dominio, las vistas son los `JFrame` y formularios `.form`, y los controladores concentran las reglas de negocio y las validaciones. La ejecucion inicia en `packagee.Main`, no directamente desde una vista.",
        "",
        "La aplicacion maneja tres tipos de usuario: administrador, paciente y doctor. El login decide a que vista entra el usuario. El administrador puede registrar doctores y abrir vistas de pacientes o doctores enviando el id correspondiente. Los pacientes pueden registrarse, actualizar su perfil, solicitar/cancelar citas y solicitar hospitalizaciones. Los doctores pueden aceptar, completar, reprogramar citas, prescribir medicamentos y resolver hospitalizaciones.",
        "",
        "La informacion inicial de usuarios se carga desde `json/users.json` usando la libreria `org.json`, incluida en `lib/json-20250107.jar`. Las citas y hospitalizaciones se simulan en memoria mediante `HospitalStore`, tal como permite el enunciado."
    ], level=2)

    add_md(lines, "2. Arquitectura MVC aplicada", [
        "**Modelos:** `User`, `Patient`, `Doctor`, `Administrator`, `Appointment`, `Hospitalization`, `Prescription` y los enums. Estas clases guardan estado y relaciones del dominio.",
        "",
        "**Vistas:** `LoginView`, `AdminView`, `PatientView` y `DoctorView`. Las vistas contienen componentes Swing, leen datos de campos, invocan controladores y muestran mensajes con `JOptionPane`. La validacion importante no se hace en la vista.",
        "",
        "**Controladores:** `UserController`, `AppointmentController`, `HospitalizationController` y `CatalogController`. Son la capa que decide si una operacion es valida, modifica modelos en `HospitalStore` y devuelve un `Response` serializado.",
        "",
        "**Almacenamiento simulado:** `HospitalStore` es un singleton que carga usuarios desde JSON, conserva listas en memoria, busca objetos por id/username y serializa modelos a JSON para que las vistas no reciban objetos de modelo directamente.",
        "",
        "**Respuesta estandar:** toda operacion importante retorna `Response`, que contiene `StatusCode`, `message` y `data`. La vista usa `message` para notificar al usuario y `data` cuando necesita pintar combos/tablas."
    ], level=2)

    add_md(lines, "3. Principios SOLID presentes", [
        "**S - Single Responsibility:** los modelos no hacen validaciones de interfaz; los controladores validan operaciones; las vistas muestran y recogen datos; `HospitalStore` administra almacenamiento y serializacion.",
        "",
        "**O - Open/Closed:** agregar nuevos estados, especialidades o tipos de respuesta se puede hacer extendiendo enums/controladores sin reescribir las vistas completas.",
        "",
        "**L - Liskov Substitution:** `Administrator`, `Patient` y `Doctor` heredan de `User`, por lo que se pueden almacenar juntos en `ArrayList<User>` y luego distinguirse con `instanceof` cuando se necesita comportamiento especifico.",
        "",
        "**I - Interface Segregation:** `StoreObserver` es una interfaz pequena con un solo metodo: `onStoreChanged`. Las vistas que necesitan refrescarse la implementan sin estar obligadas a mas metodos.",
        "",
        "**D - Dependency Inversion:** las vistas dependen de controladores y de respuestas serializadas, no de la manipulacion directa de listas de modelos. Los controladores usan `HospitalStore` como fuente comun de datos."
    ], level=2)

    class_section(lines, "Main.java",
        "Punto de entrada oficial de la aplicacion.",
        [
            "`Main` contiene el metodo `main`. Su trabajo es configurar el look and feel FlatLaf y abrir `LoginView` dentro del hilo de eventos de Swing (`EventQueue.invokeLater`). Esto cumple el requisito de que la vista no se ejecute a si misma.",
            "",
            "Primero desactiva el uso de libreria nativa de FlatLaf con `System.setProperty(\"flatlaf.useNativeLibrary\", \"false\")`. Luego intenta aplicar `FlatDarkLaf`. Si falla, no detiene el programa; imprime un mensaje de error. Finalmente crea `new LoginView().setVisible(true)`.",
        ],
        notes=[
            "Si preguntan por donde se ejecuta el proyecto: se ejecuta desde `packagee.Main`, configurado tambien en `nbproject/project.properties` como `main.class=packagee.Main`.",
        ])

    class_section(lines, "Response.java",
        "Objeto estandar de respuesta entre controladores y vistas.",
        [
            "Esta clase evita que cada controlador retorne datos de forma diferente. Tiene tres campos finales: `statusCode`, `message` y `data`.",
            "",
            "`statusCode` indica el resultado tecnico: OK, CREATED, BAD_REQUEST, NOT_FOUND, CONFLICT o UNAUTHORIZED. `message` es el texto que la vista muestra al usuario. `data` es un `String` con JSON serializado, nunca un objeto de modelo.",
            "",
            "`ok`, `created` y `error` son metodos fabrica. Esto hace mas legible el codigo del controlador: `Response.ok(...)` para exito, `Response.created(...)` cuando se crea un registro, y `Response.error(...)` para fallos. `isSuccess` concentra la regla de que solo OK y CREATED se consideran exitosos."
        ],
        notes=[
            "Respuesta clave para el PDF: cuando se retorna informacion a la vista, se retorna serializada en `data`, no como `Patient`, `Doctor` o `Appointment`."
        ])

    class_section(lines, "StatusCode.java",
        "Enum de codigos de estado usados por `Response`.",
        [
            "Define los estados de respuesta que pueden devolver los controladores. `OK` significa operacion exitosa. `CREATED` significa que se creo algo nuevo. `BAD_REQUEST` se usa para datos mal formados. `NOT_FOUND` se usa cuando no existe un usuario, cita u hospitalizacion. `CONFLICT` se usa cuando los datos son validos en forma pero chocan con una regla de negocio. `UNAUTHORIZED` se usa para login incorrecto."
        ])

    class_section(lines, "Validation.java",
        "Utilidad centralizada para validaciones de formato.",
        [
            "`Validation` tiene constructor privado para que no se instancie. Todos sus metodos son `static` porque representa una caja de herramientas de validacion.",
            "",
            "`isValidUserId` valida que el id sea mayor que 0 y tenga 12 digitos. `isValidPhone` valida 10 digitos. `isValidEmail` usa regex para exigir formato con `@` y terminacion `.com`. `isValidLicence` exige `L-XXXXXXXXXX MTL`. `isValidOffice` exige `O-XXX`.",
            "",
            "`parseDate` intenta convertir un texto a `LocalDate`; si falla retorna `null`. `parseQuarterHour` intenta convertir texto a `LocalTime` y ademas exige minutos 00, 15, 30 o 45. Esta funcion soporta la regla de que las citas duran 15 minutos y empiezan en cuartos de hora."
        ],
        notes=[
            "La vista no decide si el telefono, email, fecha, licencia u oficina son validos. La vista manda textos al controlador; el controlador llama esta clase."
        ])

    class_section(lines, "HospitalStore.java",
        "Almacenamiento simulado, carga JSON, busqueda, serializacion y observer.",
        [
            "`HospitalStore` es una clase singleton. Tiene una instancia unica en `INSTANCE`, constructor privado y metodo `getInstance`. Asi todos los controladores y vistas trabajan sobre las mismas listas en memoria.",
            "",
            "Guarda `users`, `appointments`, `hospitalizations` y `observers`. `users` se carga desde `json/users.json`. `appointments` y `hospitalizations` se llenan durante la ejecucion, simulando almacenamiento. `loaded` evita cargar el JSON mas de una vez.",
            "",
            "`loadUsersFromJson` abre `json/users.json`, lee el arreglo `users` con `JSONTokener` y por cada objeto llama `addUserFromJson`. Si ocurre un error, crea un administrador por defecto para no dejar el sistema inutilizable.",
            "",
            "`addUserFromJson` interpreta el campo `type`: si es `admin` crea `Administrator`; si es `patient` crea `Patient`; si es `doctor` crea `Doctor`. Para doctores convierte especialidad de texto a enum con `parseSpecialty`.",
            "",
            "`parseSpecialty` permite que valores como `ORTHOPEDICS` y `GYNECOLOGY` del JSON se mapeen a enums internos mas largos: `TRAUMATOLOGY_ORTHOPEDICS` y `GYNECOLOGY_OBSTETRICS`. `displaySpecialty` hace el proceso contrario para mostrar texto legible en combos.",
            "",
            "`findUserByUsername`, `findUserById`, `findPatient`, `findDoctor`, `findAppointment` y `findHospitalization` son metodos de busqueda centralizada.",
            "",
            "`serializeUser`, `serializeAppointment` y `serializeHospitalization` convierten modelos a `JSONObject`. Esto es importante porque los controladores devuelven JSON en `Response.data`. En citas, se serializan `patientId`, `patient` con nombre completo, `doctorId`, `doctor`, especialidad, fecha/hora, tipo y estado.",
            "",
            "El patron observer se implementa con `addObserver`, `removeObserver` y `notifyObservers`. Cuando un controlador modifica datos, llama `notifyObservers`; las vistas registradas refrescan combos y tablas."
        ],
        notes=[
            "Si preguntan por persistencia: no se escribe de vuelta al JSON. El enunciado permite simular almacenamiento; por eso las creaciones viven en memoria durante la ejecucion."
        ])

    class_section(lines, "StoreObserver.java",
        "Interfaz simple para refrescar vistas cuando cambia el almacenamiento.",
        [
            "Solo declara `void onStoreChanged()`. `PatientView` y `DoctorView` la implementan. Cuando se crea, acepta, cancela, completa o reprograma algo, el controlador llama `HospitalStore.notifyObservers()`, y cada vista registrada refresca sus datos."
        ])

    class_section(lines, "User.java",
        "Clase abstracta base para todos los usuarios.",
        [
            "`User` contiene los datos comunes: `id`, `username`, `firstname`, `lastname` y `password`. El `id` es `final`, por lo que no tiene setter y no se puede modificar despues de construir el usuario. Esto cumple la regla de que los ids no pueden modificarse.",
            "",
            "El resto de campos si tienen setters porque el PDF permite modificar username y datos personales. `Administrator`, `Patient` y `Doctor` heredan estos campos."
        ])

    class_section(lines, "Administrator.java",
        "Modelo del administrador.",
        [
            "Es la subclase mas simple de `User`. No agrega campos propios porque el administrador solo necesita identidad, username y password para entrar y gestionar doctores/pacientes."
        ])

    class_section(lines, "Patient.java",
        "Modelo de paciente.",
        [
            "Extiende `User` y agrega `email`, `birthdate`, `gender`, `phone`, `address`, una lista de `appointments` y una referencia a `hospitalization`.",
            "",
            "El constructor inicializa los datos personales y crea la lista de citas vacia. `addAppointment` agrega una cita a la lista del paciente. `setHospitalization` enlaza al paciente con una hospitalizacion actual o ultima asignada.",
            "",
            "Las reglas de formato de paciente no estan dentro de este modelo. Se validan en `UserController`, usando `Validation`. Esto mantiene el modelo como representacion de estado y evita mezclarlo con logica de interfaz."
        ])

    class_section(lines, "Doctor.java",
        "Modelo de doctor.",
        [
            "Extiende `User` y agrega `specialty`, `licenceNumber`, `assignedOffice`, lista de citas y lista de hospitalizaciones.",
            "",
            "El constructor configura especialidad, licencia y oficina, y crea listas vacias. `addAppointment` y `addHospitalization` mantienen relaciones desde el doctor hacia los eventos clinicos. Los setters permiten actualizar perfil del doctor, pero el id heredado sigue siendo inmutable."
        ])

    class_section(lines, "Appointment.java",
        "Modelo de cita medica.",
        [
            "Representa una cita entre paciente y doctor. Tiene `id`, `patient`, `doctor`, `specialty`, `datetime`, `reason`, `type`, lista de `prescriptions`, `status`, `diagnosis`, `observations`, `recommendedTreatment` y `followUp`.",
            "",
            "El constructor fija el estado inicial como `AppointmentStatus.REQUESTED`. Tambien agrega la cita a la lista del paciente y a la lista del doctor con `patient.addAppointment(this)` y `doctor.addAppointment(this)`. Asi se crean relaciones bidireccionales basicas.",
            "",
            "`appendReason` agrega al motivo original una razon de reprogramacion. Esto cumple la regla de que, al reprogramar, la razon nueva se debe anadir a la razon original. `addPrescription` agrega medicamentos a la cita."
        ])

    class_section(lines, "AppointmentStatus.java",
        "Enum de estados de cita.",
        [
            "`REQUESTED`: la cita fue solicitada por paciente y espera aceptacion.",
            "",
            "`PENDING`: el doctor acepto la cita y esta pendiente por completarse.",
            "",
            "`COMPLETED`: la cita fue terminada por el doctor o se completo al enviar al paciente a hospitalizacion directa.",
            "",
            "`CANCELED`: el paciente cancelo una cita que no estaba completada."
        ])

    class_section(lines, "Prescription.java",
        "Modelo de medicamento prescrito durante una cita.",
        [
            "Guarda la cita asociada, nombre del medicamento, dosis, via de administracion, duracion del tratamiento, instrucciones adicionales y frecuencia.",
            "",
            "No valida si la cita permite prescribir. Esa regla esta en `AppointmentController.prescribe`, donde solo se permite prescribir si la cita esta en estado `PENDING`."
        ],
        notes=[
            "El archivo se llama `Prescription.java`, igual que la clase publica `Prescription`, para cumplir la regla de Java."
        ])

    class_section(lines, "Hospitalization.java",
        "Modelo de hospitalizacion.",
        [
            "Representa una hospitalizacion con `id`, `patient`, `doctor`, `date`, `reason`, `roomType`, `observations` y `status`.",
            "",
            "Tiene dos constructores. El primero crea hospitalizaciones normales en estado `REQUESTED`. El segundo recibe explicitamente un `HospitalizationStatus`; se usa cuando el doctor envia al paciente directamente desde una cita y la hospitalizacion debe iniciar en `ONGOING`.",
            "",
            "En ambos constructores se enlaza la hospitalizacion con paciente y doctor: `patient.setHospitalization(this)` y `doctor.addHospitalization(this)`."
        ])

    class_section(lines, "HospitalizationStatus.java",
        "Enum de estados de hospitalizacion.",
        [
            "`REQUESTED`: hospitalizacion solicitada y pendiente de decision.",
            "",
            "`ONGOING`: hospitalizacion aprobada o creada directamente desde una cita.",
            "",
            "`CANCELED`: hospitalizacion denegada/cancelada."
        ])

    class_section(lines, "RoomType.java",
        "Enum de tipos de habitacion.",
        [
            "Define opciones de habitacion: `STANDARD`, `ICU`, `NICU`, `IMC` e `ISOLATION`. `PatientView` carga estos valores automaticamente en el combo de tipo de habitacion."
        ])

    class_section(lines, "Specialty.java",
        "Enum de especialidades medicas.",
        [
            "Centraliza las especialidades que pueden tener los doctores y las citas: medicina general, cardiologia, pediatria, neurologia, traumatologia/ortopedia, ginecologia/obstetricia, dermatologia, psiquiatria, oncologia, oftalmologia y medicina interna.",
            "",
            "Los controladores y combos trabajan con este enum para no depender de textos sueltos."
        ])

    class_section(lines, "UserController.java",
        "Controlador de login, registro y actualizacion de usuarios.",
        [
            "Al crearse, obtiene `HospitalStore.getInstance()` y asegura la carga del JSON con `store.loadUsersFromJson()`.",
            "",
            "`login(username, password)` busca por username. Si no existe o la contrasena no coincide, devuelve `UNAUTHORIZED`. Si coincide, devuelve `OK` con el usuario serializado. La vista lee el campo `type` del JSON para navegar a admin, doctor o paciente.",
            "",
            "`registerPatient` existe en version con tipos ya convertidos y version con strings. La version de strings se usa desde Swing porque los campos visuales entregan texto. Convierte `id` y `phone`; si fallan, retorna `BAD_REQUEST`. Luego llama la version tipada, que valida reglas y crea `Patient`.",
            "",
            "`updatePatient` busca el paciente por id. Si no existe retorna `NOT_FOUND`. Luego aplica las mismas reglas de creacion: id valido, username unico, password coincidente, telefono, email y fecha validos. Si todo pasa, actualiza campos permitidos y notifica observers.",
            "",
            "`registerDoctor` y `updateDoctor` hacen lo mismo para doctores, validando especialidad, licencia y oficina. El administrador usa registro de doctores; el doctor o administrador puede actualizar datos del doctor.",
            "",
            "`validateCommonUser` concentra reglas compartidas: id de 12 digitos, password igual a confirmacion, id unico y username unico. `currentUser` permite actualizar un usuario sin que choque consigo mismo.",
            "",
            "`validatePatient` agrega telefono, email y fecha de nacimiento. `validateDoctor` agrega licencia y oficina.",
        ],
        notes=[
            "Punto fuerte para explicar: la vista no valida formatos; el controlador recibe datos y decide si responde exito o error."
        ])

    class_section(lines, "CatalogController.java",
        "Controlador de lectura de datos para combos, tablas y carga de usuario.",
        [
            "`getPatients` retorna un arreglo JSON con pacientes serializados. `getDoctors` retorna doctores. `getUser(id)` busca un usuario por id y retorna su JSON. `getSpecialties` recorre el enum `Specialty` y devuelve `name` tecnico y `label` legible.",
            "",
            "Este controlador evita que las vistas entren directamente a `HospitalStore.getUsers()` para armar combos. Tambien mantiene la regla de enviar informacion serializada."
        ])

    class_section(lines, "AppointmentController.java",
        "Controlador de citas: solicitar, aceptar, completar, cancelar, reprogramar, prescribir y consultar.",
        [
            "`requestAppointment` recibe paciente, doctor opcional, especialidad, fecha, hora, razon y tipo. Primero verifica que el paciente exista. Luego valida fecha y hora con `Validation`. Despues resuelve un doctor disponible.",
            "",
            "Si el paciente selecciona doctor especifico, `resolveAvailableDoctor` busca ese doctor y verifica disponibilidad. Si el paciente selecciona especialidad, recorre usuarios doctores buscando uno con esa especialidad y disponibilidad. Si no hay doctor disponible, retorna `CONFLICT`.",
            "",
            "`nextAppointmentId` genera ids como `A-{id_paciente}-NNNN`. Cuenta citas existentes con el mismo prefijo y arma el consecutivo con cuatro digitos.",
            "",
            "`acceptAppointment` exige que la cita exista y este en `REQUESTED`. Si pasa, cambia estado a `PENDING`.",
            "",
            "`completeAppointment` exige que la cita exista y este en `PENDING`. Si pasa, cambia estado a `COMPLETED` y guarda diagnostico, observaciones, tratamiento recomendado y seguimiento.",
            "",
            "`cancelAppointment` permite cancelar citas que no esten `COMPLETED`; cambia estado a `CANCELED`.",
            "",
            "`rescheduleAppointment` no cambia el dia: toma la fecha original de la cita y solo reemplaza la hora. Valida que la nueva hora este en formato hh:mm y minutos de cuarto de hora. Tambien revisa disponibilidad del doctor y agrega la razon con `appendReason`.",
            "",
            "`prescribe` solo permite prescribir si la cita esta `PENDING`. Crea un `Prescription`, lo agrega a la cita y retorna la cita serializada.",
            "",
            "`getPatientAppointments` y `getDoctorAppointments` devuelven arreglos JSON ordenados descendentemente por fecha/hora. `getDoctorAppointments` recibe `onlyPending` para responder a la vista de doctor cuando quiere ver todas o solo pendientes."
        ],
        notes=[
            "Este controlador concentra la mayoria de reglas del PDF sobre citas y disponibilidad."
        ])

    class_section(lines, "HospitalizationController.java",
        "Controlador de hospitalizaciones.",
        [
            "`requestHospitalization` valida que paciente y doctor existan, valida fecha y crea una hospitalizacion en estado `REQUESTED`. Su id se genera con `nextHospitalizationId`, formato `H-{id_paciente}-NNNN`.",
            "",
            "`approveHospitalization` busca la hospitalizacion y cambia estado a `ONGOING`. `denyHospitalization` cambia estado a `CANCELED`. `cancelHospitalization` reutiliza `denyHospitalization`.",
            "",
            "`sendToHospitalizationFromAppointment` implementa el flujo especial del PDF: desde una cita aceptada (`PENDING`), el doctor envia al paciente a hospitalizacion. El metodo valida que la cita exista, que este `PENDING` y que la fecha sea valida; luego marca la cita como `COMPLETED` y crea hospitalizacion con estado inicial `ONGOING`.",
            "",
            "`getHospitalizations` serializa todas las hospitalizaciones para que la vista de doctor pueda cargar solicitudes pendientes en combos."
        ])

    class_section(lines, "PanelRound.java",
        "Componente Swing visual para paneles redondeados.",
        [
            "Extiende `JPanel` y sobreescribe `paintComponent` para dibujar un rectangulo redondeado con antialiasing. Es parte visual original del proyecto. No participa en reglas de negocio.",
            "",
            "`setRadius` cambia el radio y llama `repaint`. El constructor llama `setOpaque(false)` para que el fondo redondeado se pinte correctamente."
        ])

    add_md(lines, "4. Vistas / Frames", [
        "Las vistas contienen bastante codigo generado por NetBeans en `initComponents`. Ese codigo crea componentes, asigna fuentes, textos, layouts y eventos. No conviene editarlo manualmente salvo con cuidado, porque NetBeans lo usa para abrir en modo Design. La logica propia se ubica en constructores, handlers `jButton...ActionPerformed`, metodos `load...`, `refresh...` y `configureComponentNames`.",
        "",
        "Durante el refactor se conservaron posiciones, componentes y aspecto visual. Lo que se cambio fue el nombre de las clases y variables para que fueran claras, y se movieron validaciones a controladores."
    ], level=2)

    class_section(lines, "LoginView.java",
        "Frame de login y registro de pacientes.",
        [
            "Antes era `NewJFrame`. Ahora se llama `LoginView`, nombre que describe su funcion. El constructor llama `initComponents`, carga usuarios desde JSON y luego ejecuta `configureComponentNames`.",
            "",
            "La vista tiene dos areas principales: login y registro de paciente. En login aparecen usuario, contrasena y boton ENTER. En registro aparecen campos de paciente: nombre, apellido, id, genero, telefono, email, usuario, password, confirmacion, direccion y fecha de nacimiento.",
            "",
            "`jButton2ActionPerformed` es el login. Crea `UserController`, llama `login`, muestra `response.getMessage()` y, si es exitoso, interpreta `response.getData()` como JSON. Segun `type`, navega a `AdminView`, `DoctorView` o `PatientView`. En todos los casos pasa ids, no objetos completos.",
            "",
            "`jButton9ActionPerformed` registra paciente. La vista toma textos de campos y seleccion de genero, llama `UserController.registerPatient`, muestra mensaje y si la respuesta es exitosa limpia los campos. Las reglas de id, telefono, email, fecha y password se validan en el controlador.",
            "",
            "`jButton1ActionPerformed` cierra la aplicacion con `System.exit(0)`. Los eventos `panelRound2MousePressed` y `panelRound2MouseDragged` permiten arrastrar la ventana por la barra superior.",
            "",
            "`configureComponentNames` asigna nombres semanticos a botones, campos y tabs, por ejemplo `loginSubmitButton`, `patientRegisterIdField` y `loginAndRegistrationTabs`."
        ],
        notes=[
            "Explicacion simple: `LoginView` no decide si las credenciales son correctas; pregunta al controlador y actua segun la respuesta."
        ])

    class_section(lines, "AdminView.java",
        "Frame del administrador.",
        [
            "Antes era `NewJFrame11`. Ahora se llama `AdminView`. El constructor recibe `userId`, carga usuarios, llena combos de doctores/pacientes y configura nombres de componentes.",
            "",
            "Visualmente permite registrar doctores y abrir vista de doctor o paciente. El administrador no necesita modificar el diseno; solo usa botones ya existentes.",
            "",
            "`jButton9ActionPerformed` registra doctores. Recoge id, username, nombres, password, especialidad, licencia y oficina; llama `UserController.registerDoctor`; muestra respuesta; si es exitoso recarga combos y limpia campos.",
            "",
            "`jButton2ActionPerformed` abre `DoctorView` como administrador. Toma el id del combo de doctores, valida con `CatalogController.getUser`, verifica que el tipo sea `doctor`, y crea `new DoctorView(userId, doctorData.getLong(\"id\"))`. Asi el doctor mostrado puede ser diferente del usuario logueado.",
            "",
            "`jButton3ActionPerformed` abre `PatientView` como administrador con el mismo patron, verificando tipo `patient`.",
            "",
            "`jButton10ActionPerformed` hace logout: abre `LoginView` y oculta la vista actual.",
            "",
            "`loadUsersInCombos` consulta `CatalogController.getDoctors` y `getPatients`, y llena combos con ids. `configureComponentNames` renombra todos los componentes graficos importantes."
        ],
        notes=[
            "Punto del PDF cumplido aqui: el administrador puede acceder a vistas de paciente y doctor enviando la informacion correspondiente."
        ])

    class_section(lines, "PatientView.java",
        "Frame del paciente.",
        [
            "Antes era `NewJFrame1`. Ahora se llama `PatientView`. Implementa `StoreObserver`, por lo que puede refrescarse cuando cambian datos del store. El constructor recibe `userId` y `patientId`: `userId` indica quien entro; `patientId` indica que paciente se esta visualizando.",
            "",
            "Si el usuario real es admin, el boton Back se muestra. Si el usuario real no es admin, se oculta. Esto cumple que el boton back solo este activo para administradores.",
            "",
            "La vista tiene pestañas/secciones para historial de citas, perfil y solicitudes. `loadPatientData` consulta `CatalogController.getUser(patientId)` y carga nombre, apellido, fecha, genero, email, telefono, direccion y username en los campos visuales.",
            "",
            "`jButton9ActionPerformed` actualiza el perfil del paciente. Recoge datos, llama `UserController.updatePatient`, muestra mensaje y en exito recarga datos.",
            "",
            "`jRadioButton3ActionPerformed` activa modo solicitud por especialidad y llama `loadSpecialtiesCombo`. `jRadioButton4ActionPerformed` activa modo solicitud por doctor y llama `loadDoctorsCombo`.",
            "",
            "`jButton3ActionPerformed` solicita cita. Si esta seleccionado doctor, manda doctorId y especialidad vacia; si esta seleccionado especialidad, manda especialidad y doctor vacio. Luego llama `AppointmentController.requestAppointment`. En exito limpia fecha, hora y razon, refresca tabla y combo de cancelacion.",
            "",
            "`jButton4ActionPerformed` solicita hospitalizacion como paciente. Usa `HospitalizationController.requestHospitalization`, pasando paciente actual, doctor seleccionado, fecha, motivo, tipo de habitacion y observaciones. Si sale bien limpia campos.",
            "",
            "`jButton5ActionPerformed` cancela una cita seleccionada con `AppointmentController.cancelAppointment`.",
            "",
            "`refreshAppointmentsTable` consulta `AppointmentController.getPatientAppointments` y llena la tabla con id, fecha/hora, doctor, especialidad, tipo y estado. `loadPatientAppointmentCombo` carga solo citas no completadas para cancelar.",
            "",
            "`onStoreChanged` recarga combos y tabla cuando el store notifica cambios."
        ],
        notes=[
            "La vista muestra respuestas, limpia campos en exito y no valida formatos por su cuenta."
        ])

    class_section(lines, "DoctorView.java",
        "Frame del doctor.",
        [
            "Antes era `NewJFrame111`. Ahora se llama `DoctorView`. Implementa `StoreObserver`. El constructor recibe `userId` y `doctorId`, carga el doctor, combos, tabla de citas y registra la vista como observer.",
            "",
            "Si el usuario real es admin, muestra el boton Back; si es doctor normal, lo oculta. Esto permite que el administrador vuelva a `AdminView` desde vista doctor.",
            "",
            "La vista tiene pestañas para citas del doctor, historial de paciente, perfil, acciones sobre citas/hospitalizaciones y prescripcion de medicamentos.",
            "",
            "`refreshDoctorAppointments(onlyPending)` consulta `AppointmentController.getDoctorAppointments`. Si `onlyPending` es true, muestra pendientes; si es false, muestra todas. La tabla muestra id, fecha, nombre completo del paciente, especialidad, tipo y estado.",
            "",
            "`jRadioButton3ActionPerformed` y `jRadioButton4ActionPerformed` alternan entre todas las citas y citas pendientes.",
            "",
            "`jButton3ActionPerformed` acepta una cita seleccionada: llama `AppointmentController.acceptAppointment`. El controlador verifica que este `REQUESTED` y la pasa a `PENDING`.",
            "",
            "`jButton4ActionPerformed` reprograma una cita: llama `AppointmentController.rescheduleAppointment` con id, nueva hora y razon. El controlador no permite cambiar fecha, solo hora, y agrega la razon al motivo original.",
            "",
            "`jButton5ActionPerformed` completa una cita: llama `AppointmentController.completeAppointment` con diagnostico, observaciones, tratamiento y seguimiento. El controlador exige estado `PENDING`.",
            "",
            "`jButton7ActionPerformed` prescribe medicamento: toma cita, medicamento, dosis, via, duracion, instrucciones y frecuencia; llama `AppointmentController.prescribe`; si sale bien agrega fila a la tabla de prescripciones y limpia campos.",
            "",
            "`jButton6ActionPerformed` genera hospitalizacion. Si hay una cita seleccionada en el combo de citas completables, usa `HospitalizationController.sendToHospitalizationFromAppointment`, lo que completa la cita y crea hospitalizacion `ONGOING`. Si no hay cita seleccionada, usa flujo normal: `requestHospitalization` y luego `approveHospitalization` para dejarla `ONGOING`.",
            "",
            "`jButton13ActionPerformed` cancela/deniega solicitudes de hospitalizacion seleccionadas con `HospitalizationController.denyHospitalization`.",
            "",
            "`jButton8ActionPerformed` consulta historial de un paciente con `AppointmentController.getPatientAppointments` y llena la tabla de historial.",
            "",
            "`loadAppointmentCombos` llena combos segun estados: citas `REQUESTED` para aceptar, citas no canceladas ni completadas para reprogramar, citas `PENDING` para completar y prescribir.",
            "",
            "`loadHospitalizationCombo` carga solicitudes `REQUESTED` asignadas al doctor. `onStoreChanged` recarga combos y tabla cuando cambia el store."
        ],
        notes=[
            "Aqui estan los flujos mas preguntables en sustentacion: aceptar, completar, reprogramar, prescribir y hospitalizacion directa desde cita."
        ])

    add_md(lines, "5. Archivos .form de NetBeans", [
        "`LoginView.form`, `AdminView.form`, `PatientView.form` y `DoctorView.form` guardan la metadata visual de NetBeans GUI Builder. Estos archivos deben mantenerse consistentes con los `.java` para que cada frame abra en modo Design.",
        "",
        "No se deben editar como texto salvo que se sepa exactamente que se hace. En este proyecto se mantuvo el aspecto visual; el cambio principal fue renombrar clases y componentes, conservando estructura y ubicacion visual."
    ], level=2)

    add_md(lines, "6. JSON inicial", [
        "`json/users.json` contiene usuarios iniciales. Tiene administradores, pacientes y doctores. Cada objeto tiene `type`, `id`, `username`, `firstname`, `lastname` y `password`. Los pacientes tienen ademas `email`, `birthdate`, `gender`, `phone` y `address`. Los doctores tienen `specialty`, `licenceNumber` y `assignedOffice`.",
        "",
        "Credenciales utiles para probar:",
        "",
        "- Admin: `admin_root` / `Admin@1234`",
        "- Paciente: `jgarcia90` / `Pass@1234`",
        "- Doctor: `dr_aguirre` / `Doc@1234`"
    ], level=2)

    add_md(lines, "7. Build y librerias", [
        "`nbproject/project.properties` define `main.class=packagee.Main`, por eso Run debe iniciar desde `Main`. El classpath incluye `lib/flatlaf-demo-3.6.jar` para apariencia FlatLaf y `lib/json-20250107.jar` para `org.json`.",
        "",
        "`build.xml` es el build Ant de NetBeans. Importa `nbproject/build-impl.xml`, por lo que NetBeans puede compilar y ejecutar el proyecto con su configuracion normal."
    ], level=2)

    add_md(lines, "8. Flujos principales explicados paso a paso", level=2)
    flows = [
        ("Login admin", ["Usuario escribe credenciales en `LoginView`.", "`LoginView` llama `UserController.login`.", "Controlador busca en `HospitalStore`.", "Si el JSON retornado tiene `type=admin`, se abre `AdminView`."]),
        ("Login paciente", ["Mismo login.", "Si `type=patient`, se abre `PatientView(userId, patientId)`.", "La vista carga datos automaticamente con `CatalogController.getUser`."]),
        ("Login doctor", ["Mismo login.", "Si `type=doctor`, se abre `DoctorView(userId, doctorId)`.", "La vista carga citas, perfil, combos y hospitalizaciones."]),
        ("Registro paciente", ["`LoginView` toma campos.", "Llama `UserController.registerPatient`.", "Controlador valida id, username, password, telefono, email y fecha.", "Crea `Patient`, notifica observers y retorna JSON."]),
        ("Registro doctor", ["`AdminView` toma campos de doctor.", "Llama `UserController.registerDoctor`.", "Controlador valida id, username, password, licencia y oficina.", "Crea `Doctor`, notifica observers y retorna JSON."]),
        ("Solicitud de cita", ["`PatientView` manda paciente, doctor opcional o especialidad, fecha, hora, razon y tipo.", "`AppointmentController` valida paciente, fecha, hora y disponibilidad.", "Genera id `A-{paciente}-NNNN`.", "Crea `Appointment` en estado `REQUESTED`."]),
        ("Aceptar cita", ["`DoctorView` carga citas `REQUESTED` en combo.", "Doctor selecciona una.", "`AppointmentController.acceptAppointment` valida estado `REQUESTED`.", "Estado pasa a `PENDING`."]),
        ("Completar cita", ["`DoctorView` carga citas `PENDING`.", "Doctor ingresa diagnostico/observaciones/tratamiento/seguimiento.", "`AppointmentController.completeAppointment` valida estado `PENDING`.", "Estado pasa a `COMPLETED`."]),
        ("Reprogramar cita", ["Doctor selecciona cita no cancelada ni completada.", "Ingresa nueva hora y razon.", "Controlador valida hora en cuartos y disponibilidad.", "Actualiza hora sin cambiar fecha y agrega razon."]),
        ("Prescribir", ["Doctor selecciona cita `PENDING`.", "Ingresa medicamento, dosis, via, duracion, instrucciones y frecuencia.", "Controlador valida numeros y estado.", "Agrega `Prescription` a la cita."]),
        ("Hospitalizacion paciente", ["Paciente selecciona doctor, fecha, habitacion, razon y observaciones.", "Controlador valida paciente/doctor/fecha.", "Genera id `H-{paciente}-NNNN` en estado `REQUESTED`."]),
        ("Hospitalizacion directa desde cita", ["Doctor selecciona cita `PENDING` en combo de completar.", "En seccion hospitalizacion presiona Generate.", "`DoctorView` llama `sendToHospitalizationFromAppointment`.", "Controlador marca cita `COMPLETED` y crea hospitalizacion `ONGOING`."]),
    ]
    for title, steps in flows:
        add_md(lines, title, level=3)
        for i, step in enumerate(steps, 1):
            lines.append(f"{i}. {step}")
        lines.append("")

    add_md(lines, "9. Reglas del PDF y donde viven", [
        "- IDs de usuarios: `UserController.validateCommonUser` y `Validation.isValidUserId`.",
        "- Username unico: `UserController.validateCommonUser` con `HospitalStore.findUserByUsername`.",
        "- Password y confirmacion: `UserController.validateCommonUser`.",
        "- Telefono paciente: `Validation.isValidPhone`.",
        "- Email paciente: `Validation.isValidEmail`.",
        "- Fecha paciente/cita/hospitalizacion: `Validation.parseDate`.",
        "- Licencia doctor: `Validation.isValidLicence`.",
        "- Oficina doctor: `Validation.isValidOffice`.",
        "- Hora cita: `Validation.parseQuarterHour`.",
        "- Doctor valido/disponible: `AppointmentController.resolveAvailableDoctor` e `isDoctorAvailable`.",
        "- Estado inicial cita: constructor de `Appointment`.",
        "- Aceptar/completar/cancelar/reprogramar: `AppointmentController`.",
        "- Prescribir solo cita aceptada: `AppointmentController.prescribe`.",
        "- ID cita automatico: `AppointmentController.nextAppointmentId`.",
        "- ID hospitalizacion automatico: `HospitalizationController.nextHospitalizationId`.",
        "- Hospitalizacion directa desde cita: `HospitalizationController.sendToHospitalizationFromAppointment`.",
        "- Orden descendente de citas: `AppointmentController.sortedAppointments`.",
        "- No retornar modelos a vista: `HospitalStore.serialize...` + `Response.data`."
    ], level=2)

    add_md(lines, "10. Detalles importantes para no confundirse en sustentacion", [
        "1. El proyecto no guarda cambios nuevos en el JSON; los mantiene en memoria. Eso es una simulacion de almacenamiento.",
        "2. Los objetos si existen internamente como modelos, pero nunca se envian directamente a la vista como respuesta; se serializan a JSON.",
        "3. Las vistas si leen campos de texto, porque eso es su responsabilidad, pero no validan reglas de negocio.",
        "4. Los nombres `jButton...ActionPerformed` permanecen porque NetBeans genera esos nombres para handlers; los componentes visuales si tienen nombres claros como variables y con `setName`.",
        "5. El boton Back en paciente/doctor depende de si `userId` corresponde a admin.",
        "6. `doctorId` o `patientId` indica que perfil se esta visualizando; `userId` indica quien inicio sesion.",
        "7. `REQUESTED`, `PENDING`, `COMPLETED` y `CANCELED` son estados de cita; `REQUESTED`, `ONGOING` y `CANCELED` son estados de hospitalizacion.",
        "8. La disponibilidad de doctor considera citas del mismo doctor en la misma fecha/hora que no esten canceladas."
    ], level=2)

    add_md(lines, "11. Mapa rapido de clases", level=2)
    mapping = [
        ("Main", "Arranca la app y abre LoginView."),
        ("HospitalStore", "Singleton con datos, carga JSON, busquedas, serializacion y observer."),
        ("Validation", "Validaciones reutilizables de formatos."),
        ("Response / StatusCode", "Contrato comun de respuesta."),
        ("UserController", "Login, registro y actualizacion de usuarios."),
        ("AppointmentController", "Reglas de citas."),
        ("HospitalizationController", "Reglas de hospitalizaciones."),
        ("CatalogController", "Datos para combos/tablas/perfiles."),
        ("LoginView", "Login y registro paciente."),
        ("AdminView", "Registro doctor y navegacion admin."),
        ("PatientView", "Perfil paciente, citas, hospitalizacion, cancelacion."),
        ("DoctorView", "Citas doctor, historial, perfil, prescripcion, hospitalizacion."),
        ("User / Patient / Doctor / Administrator", "Modelos de usuarios."),
        ("Appointment / Hospitalization / Prescription", "Modelos clinicos."),
        ("Enums", "Estados, especialidades y tipos de habitacion.")
    ]
    for name, desc in mapping:
        lines.append(f"- `{name}`: {desc}")
    lines.append("")

    add_md(lines, "12. Checklist de pruebas manuales recomendadas", [
        "- Ejecutar desde NetBeans con Run y confirmar que inicia `LoginView`.",
        "- Abrir cada `.form` en Design: `LoginView.form`, `AdminView.form`, `PatientView.form`, `DoctorView.form`.",
        "- Login admin: `admin_root` / `Admin@1234`.",
        "- Login paciente: `jgarcia90` / `Pass@1234`.",
        "- Login doctor: `dr_aguirre` / `Doc@1234`.",
        "- Registrar doctor con licencia `L-1234567899 MTL` y oficina `O-999`.",
        "- Solicitar cita con fecha `2026-06-01` y hora `09:00`.",
        "- Aceptar cita como doctor y verificar estado `PENDING`.",
        "- Prescribir en cita `PENDING`.",
        "- Completar cita y verificar estado `COMPLETED`.",
        "- Intentar cancelar cita completada y confirmar que el controlador lo rechaza.",
        "- Solicitar hospitalizacion y aprobar/denegar segun flujo.",
        "- Probar hospitalizacion directa desde cita aceptada."
    ], level=2)

    return "\n".join(lines)


def add_paragraph_doc(doc, text, style=None, bold_prefix=False):
    p = doc.add_paragraph(style=style)
    if bold_prefix and ":" in text:
        prefix, rest = text.split(":", 1)
        r = p.add_run(prefix + ":")
        r.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p


def build_docx(md_text):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.color.rgb = RGBColor(31, 77, 120)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = RGBColor(31, 77, 120)

    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(20)
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(31, 77, 120)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            continue
        if line.startswith("- "):
            add_paragraph_doc(doc, line[2:], style="List Bullet")
            continue
        if re.match(r"^\d+\. ", line):
            add_paragraph_doc(doc, re.sub(r"^\d+\. ", "", line), style="List Number")
            continue
        if line.startswith("**") and "**" in line[2:]:
            clean = line.replace("**", "")
            add_paragraph_doc(doc, clean, bold_prefix=True)
            continue
        add_paragraph_doc(doc, line)

    doc.add_page_break()
    doc.add_heading("Apendice A - Inventario automatico de archivos Java", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for idx, value in enumerate(["Archivo", "Lineas", "Campos detectados", "Metodos/constructores detectados"]):
        hdr[idx].text = value
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for path in sorted(SRC.glob("*.java")):
        count, fields, methods = source_stats(path.name)
        cells = table.add_row().cells
        cells[0].text = path.name
        cells[1].text = str(count)
        cells[2].text = str(len(fields))
        cells[3].text = str(len(methods))
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)

    doc.save(OUT_DOCX)


def main():
    md = build_markdown()
    OUT_MD.write_text(md, encoding="utf-8")
    build_docx(md)
    print(OUT_MD)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
